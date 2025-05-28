import nltk
import re
import random
import spacy 
import string
from io import BytesIO
import time
import PyPDF2
from PyPDF2 import PdfReader
import docx
from nltk.corpus import wordnet as wn
from collections import Counter
from docx import Document

try:
    wn.ensure_loaded() 
except LookupError:
    print('downloading wordnet...')
    nltk.download('wordnet', quiet=True) 

nlp = spacy.load('en_core_web_lg')

def extract_text(uploaded_file):
    if uploaded_file.type == "text/plain":
        return uploaded_file.getvalue().decode("utf-8")
    elif uploaded_file.type == "application/pdf":
        pdf_reader = PyPDF2.PdfReader(uploaded_file)
        return " ".join([page.extract_text() for page in pdf_reader.pages])
    elif "wordprocessingml" in uploaded_file.type:  # DOCX
        doc = Document(uploaded_file)
        return "\n".join([para.text for para in doc.paragraphs])
    return "Unsupported file type"

def read_text_from_file(file_path):
    if file_path.endswith(".txt"):
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    elif file_path.endswith(".pdf"):
        reader = PdfReader(file_path)
        return "".join([page.extract_text() for page in reader.pages if page.extract_text()])

    elif file_path.endswith(".docx"):
        doc = docx.Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    else:
        raise ValueError("Unsupported file type")

def extract_keywords(sent):
    doc = nlp(sent)
    
    filtered_tokens = [
        token for token in doc 
        if not token.is_stop and not token.is_punct
    ]

    keywords = []
    for token in filtered_tokens:
        if token.pos_ == "NOUN":
            if wn.synsets(token.text, pos='n'):  
                return token.text
        elif token.pos_ in ("VERB", "ADJ") and wn.synsets(token.text):
            keywords.append(token.text)  
    return keywords[0] if keywords else None

def extract_distractors(word):
    doc = nlp(word)
    lemma = doc[0].lemma_.lower()
    distractors = []

    synsets = wn.synsets(lemma, pos='n')
    for syn in synsets:
        distractors += get_distractors_wordnet(syn, lemma)
        distractors = list(set(distractors))  # Remove duplicates
        distractors = [d for d in distractors if d.lower() != lemma.lower()]
        if len(distractors) >= 2:
            break  
    while len(distractors) < 2:
        distractors.append(f"dummy{len(distractors)+1}") 

    return distractors[:3]  # Return max 3

def get_distractors_wordnet(syn, word):
    distractors = []
    word = word.lower().replace(" ", "_")
    orig_word = word
    hypernyms = syn.hypernyms()
    if not hypernyms:
        return distractors
    for item in hypernyms[0].hyponyms():
        name = item.lemmas()[0].name()
        if name == orig_word:
            continue
        name = " ".join(w.capitalize() for w in name.replace("_", " ").split()).lower()
        if name and name not in distractors:
            distractors.append(name)
    return distractors

def format_mcqs(questions):
    formatted = {}

    for qid, qdata in questions.items():
        all_options = qdata['options'] + [qdata['answer']]
        random.shuffle(all_options)

        # Map options to letters A, B, C, ...
        option_labels = list(string.ascii_uppercase)
        options = {option_labels[i]: opt for i, opt in enumerate(all_options)}

        # Find which letter is the correct answer
        correct_letter = next(k for k, v in options.items() if v == qdata['answer'])

        formatted[qid] = {
            'question': qdata['question'],
            'options': options,
            'answer': correct_letter
        }

    return formatted

def text_to_questions_pipeline(text, max_choice=5):

    text = text.replace("\n", " ").strip()
    text = re.sub(r'\s+', ' ', text)
    doc = nlp(text)
    sentences = [sent.text for sent in doc.sents if sum(1 for token in sent if token.pos_ == 'NOUN' and not token.is_stop and not token.is_punct) >= 2]
    chosen_sents = random.sample(sentences, min(max_choice, len(sentences)))
    questions = {}
    for i, sent in enumerate(chosen_sents):
        keyword = extract_keywords(sent)
        if keyword:
            distractors = extract_distractors(keyword)

            question_sent = sent.replace(keyword, "_______")
            options = distractors[:3]
            random.shuffle(options)
            questions[i] = {
                'question': question_sent,
                'answer': keyword.lower(),
                'options': options
            }
    return format_mcqs(questions)


if __name__ == "__main__":
    start_time = time.time()
    file_path = "input/Lect_01.pdf"
    file_content = read_text_from_file(file_path)
    questions = text_to_questions_pipeline(file_content, max_choice=10)

    end_time = time.time()
    elapsed_time = end_time - start_time

    for q_id, q in questions.items():
        print(f"Q{q_id}: {q['question']}")
        for opt, val in q['options'].items():
            print(f"  {opt}) {val}")
        print(f"Answer: {q['answer']}\n")

    print(f"Time taken: {elapsed_time:.2f} seconds")